from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE

from docx.shared import Pt

from docx.shared import Inches, Cm

from PyQt5.QtGui import QStandardItemModel

import locale
import copy


class DocumentComposer:

    def __init__(self):
        super(DocumentComposer, self).__init__()
        self.document = Document()
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        self.data = []
        self.spanRows = []
        self.realSpanRows = []
        self.customers = []
        self.customerCount = 0

        self.table = self.document.add_table(rows=1, cols=4)
        self.table.autofit = True
        self.table.style = 'TableGrid'
        self.currentCustomer = None

        self.pageCount = 0

    def set_result_bg_color(self):
        """
        Задает задний фон для строчек с подсчетом "ИТОГО".
        """

        lastRow = self.table.rows[len(self.table.rows) - 1]
        self.make_rows_bold(lastRow)
        cells = lastRow.cells
        if cells[0].text != 'ИТОГО':
            return
        for cell in cells:
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="ff8000"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm_1)

    def set_header_bg_color(self, rowIndex=None):
        """
        Задает задний фон для строчек с заголовком (цена/кол-во и т.д.).
        """
        localIndex = len(self.table.rows) - 1
        if rowIndex is not None:
            localIndex = rowIndex
        lastRow = self.table.rows[localIndex]
        self.make_rows_bold(lastRow)
        cells = lastRow.cells
        for cell in cells:
            # НЕ РАБОТАЕТ
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="878787"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm_1)

    @staticmethod
    def make_rows_bold(*rows):
        """
        Делает строчки жирными.
        """
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    @staticmethod
    def make_rows_underline(cells):
        """
        Делает строчки подчеркнутыми.
        """
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.underline = True

    def span_rows(self):
        """
        Объединяет строчки из realSpanRows.
        (row, column, spanData)
        """
        for row, column, spanData in self.realSpanRows:
            cell = self.table.rows[row + 2].cells[column]
            cell.merge(self.table.rows[row + spanData + 1].cells[column])

    def appendHeader(self, date):
        """
        Добавление фразы с датой из QDateEdit -> приходит к нам в формате "2" Августа 2021 г.
        """
        locale.setlocale(locale.LC_ALL, "ru_RU.UTF-8")
        date = '"' + str(date.day) + '"' + " " + str(date.strftime('%B')).title() + " " + str(date.year) + " г."
        titleName = "Заявка на " + date
        cell = self.table.rows[0].cells[0]
        cell.text = titleName
        cell.merge(self.table.rows[0].cells[3])
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.make_rows_bold(self.table.rows[0])
        self.make_rows_underline(self.table.rows[0].cells)
        self.pageCount += 1

    def appendCustomer(self, customer):
        """
        Добавление нового заказчика в таблицу если такой имеется.
        """
        if customer is not None and self.currentCustomer is None:
            self.currentCustomer = customer
        elif customer == self.currentCustomer:
            return
        else:
            self.currentCustomer = customer
        self.table.add_row()
        self.pageCount += 1
        cell = self.table.rows[len(self.table.rows) - 1].cells[0]

        self.customers.append(len(self.table.rows) - 1)

        cell.text = self.currentCustomer.strip()

        cell.bold = True
        cell.underline = True
        self.make_rows_bold(self.table.rows[len(self.table.rows) - 1])
        cell.merge(self.table.rows[len(self.table.rows) - 1].cells[3])

    def appendDataToTable(self, model, customer):
        """
        Добавляет в общий контейнер с данными запись из модели.
        Для удобства все хранится в list.
        """
        newTable = []
        headers = []
        totalLen = 1

        if customer is None:
            return

        self.spanRows = []
        rowCount = 0
        rowBefore = 0

        self.customerCount += 1

        for item in self.data:
            rowBefore += len(item)

        for item in self.data:
            rowCount += len(item)

        for column in range(model.columnCount()):
            headers.append(model.horizontalHeaderItem(column).text())

        for row in range(model.rowCount()):
            newTable.append([])
            for column in range(model.columnCount()):
                index = model.index(row, column)
                headerData = model.item(row, column).data(4)
                if headerData is not None and row + rowCount + 1 not in self.spanRows:
                    self.spanRows.append(row + rowCount + 1)

                spanData = model.itemFromIndex(index).data(5)
                if spanData is not None:
                    self.realSpanRows.append((row + rowBefore + self.customerCount, column, spanData))

                newTable[row].append(str(model.data(index)))

        self.appendCustomer(customer)
        newTable.insert(0, headers)

        self.data.append(newTable)
        self.appendTableToFile(newTable, len(self.data) + totalLen - 1)

    def mergeCell(self, cell1, cell2, row):
        """
        Объединение ячеек.
        """
        cell1.merge(cell2)
        cell1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.make_rows_bold(self.table.rows[row])
        self.make_rows_underline(self.table.rows[row].cells)

    def appendTableToFile(self, table, i):
        """
        Добавляет таблицу в файл.
        """
        for row in range(len(table)):
            self.table.add_row()
            for column in range(len(table[row])):
                cell = self.table.rows[len(self.table.rows) - 1].cells[column]
                cell.text = table[row][column]

            if row == 0:
                self.set_header_bg_color()

        for rowSpan in self.spanRows:
            self.set_header_bg_color(rowSpan + i + 1)

        self.set_result_bg_color()

    def appendEndTable(self):
        """
        Добавляет в конец документа таблицу согласно требованиям заказчика.
        """
        table = self.document.add_table(rows=6, cols=2)
        table.style = 'TableGrid'
        data = (
            ('расходы', 'сумма'),
            ('Получено с кассы', ''),
            ('Оплачено ООО', ''),
            ('ГСМ/ком-ные', ''),
            ('Итого расход', ''),
            ('Сдано в кассу', '')
        )
        for row in range(6):
            cells = table.rows[row].cells
            for col in range(2):
                cells[col].text = data[row][col]
        self.make_rows_bold(table.rows[0])

    def writeTablesToFile(self):
        """
        Для удобства записи большого количества таблиц в файл.
        """
        i = 0
        for table in self.data:
            self.appendTableToFile(table, i)
            i += 1

    def saveToFile(self, name, fileFormat, date):
        """
        Общий метод сохранения.
        Содержит в себе вызовы всех вспомогательных методов класса.
        """
        self.appendHeader(date)
        self.span_rows()

        self.copyHeader()

        self.appendEndTable()
        self.document.save(str(name) + str(fileFormat))
        self.data = []
        self.spanRows = []
        self.realSpanRows = []
        self.customerCount = 0

        self.document = Document()
        self.table = self.document.add_table(rows=1, cols=4)
        self.table.autofit = True

        self.table.style = 'TableGrid'

        self.currentCustomer = None

    def copyHeader(self):
        """
        Позволяет вставить нужные названия на следующих страницах.
        """
        rowCount = len(self.table.rows)
        current_row = self.table.rows[2]
        rowWithHeader = copy.deepcopy(current_row._tr)
        copyCustomers = []
        for rowindex in self.customers:
            searchRow = self.table.rows[rowindex]
            copyCustomers.append(copy.deepcopy(searchRow._tr))
        lenPerPage = 33
        indexes = []

        while lenPerPage < rowCount:
            indexes.append(lenPerPage)
            lenPerPage += 31

        t = self.table
        for index in indexes:
            row0 = t.rows[index]
            row0._tr.addnext(rowWithHeader)
            var = list(filter(lambda x: x < index, self.customers))[-1]
            toInsert = self.customers.index(var)
            row0._tr.addnext(copyCustomers[toInsert])
